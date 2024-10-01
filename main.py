import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import spacy
from collections import Counter

# Load NLP model
nlp = spacy.load("en_core_web_sm")

# Define common stop words
STOP_WORDS = set([
    'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 'your', 
    'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', 
    'her', 'hers', 'herself', 'it', 'its', 'itself', 'they', 'them', 'their', 
    'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', 
    'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 
    'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 
    'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 
    'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 
    'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 
    'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 
    'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 
    'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 
    'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 
    'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', 
    'should', 'now'
])

# Load the document
def load_document(file_path):
    return docx.Document(file_path)

# Function to generate HTML output
def generate_html_output(overview, errors, file_name):
    # CSS Styling
    html_content = '''
    <html>
    <head>
        <style>
            body {{font-family: Arial, sans-serif; margin: 40px;}}
            .error {{color: red; font-weight: bold;}}
            .sentence {{margin: 10px 0; padding: 10px; background-color: #f9f9f9; border-radius: 5px;}}
            .error-description {{color: darkred; margin-left: 20px;}}
            h1, h2 {{color: #333;}}
            .overview {{margin-bottom: 20px; padding: 10px; border: 1px solid #ccc; background-color: #f0f8ff; border-radius: 5px;}}
            .tag {{display: inline-block; margin: 5px; padding: 5px 10px; background-color: #e0f7fa; border-radius: 5px; border: 1px solid #81d4fa;}}
            .show-more-link {{
                color: #007bff; /* Blue color */
                cursor: pointer; /* Pointer cursor for hover effect */
                text-decoration: underline; /* Underline to indicate it's clickable */
            }}
            .show-more-link:hover {{
                color: #0056b3; /* Darker blue on hover */
            }}
            li {{ margin-bottom: 5px; }} /* Add spacing between error types */
            .hidden {{display: none;}} /* Initially hide the hidden keywords */
            footer {{
                margin-top: 20px; 
                font-size: 14px; /* Increase font size */
                color: #666; 
                text-align: center; /* Center align text */
                padding: 10px 0; /* Add padding for spacing */
            }}
        </style>
        <script>
            function toggleKeywords() {{
                var hiddenKeywords = document.getElementById('hidden-keywords');
                var link = document.getElementById('show-more-link');
                if (hiddenKeywords.style.display === 'none') {{
                    hiddenKeywords.style.display = 'block';
                    link.innerText = 'Show Less';
                }} else {{
                    hiddenKeywords.style.display = 'none';
                    link.innerText = 'Show More';
                }}
            }}
        </script>
    </head>
    <body>
        <h1>Document Review Report</h1>
        <p><strong>File Name:</strong> {file_name}</p>  <!-- Display file name -->
        <div class="overview">
            <h2>Overview</h2>
            <p><strong>Total Word Count:</strong> {word_count}</p>
            <p><strong>Page Count:</strong> {page_count}</p>
            <p><strong>Key Words:</strong></p>
            <div>
                {top_key_words}
            </div>
            <div id="hidden-keywords" class="hidden">
                {remaining_key_words}
            </div>
            <span id="show-more-link" class="show-more-link" onclick="toggleKeywords()">Show More</span>
            <p><strong>Total Errors:</strong> {total_errors}</p>
            <p><strong>Error Types:</strong></p>
            <ul>
                {error_types}
            </ul>
        </div>
        <p>The following issues were detected in the document:</p>
        {error_sections}
        <footer>
            <p>Tool created by Ashwanth V</p>
            <p>Email: <a href="mailto:22f3001662@ds.study.iitm.ac.in">22f3001662@ds.study.iitm.ac.in</a></p>
        </footer>
    </body>
    </html>
    '''
    
    # Populate the content with errors
    error_sections = ""
    for sentence, error in errors:
        error_sections += f'''
        <div class="sentence">
            <p><strong>Sentence:</strong> {sentence}</p>
            <p class="error-description"><span class="error">Error:</span> {error}</p>
        </div>
        '''
    
    # Prepare error types for the overview
    error_counter = Counter(error for _, error in errors)
    error_types = "".join(f"<li>{error}: {count}</li>" for error, count in error_counter.items())

    # Prepare key words styled as tags
    sorted_key_words = sorted(overview['key_words'].items(), key=lambda x: x[1], reverse=True)
    top_key_words = "".join(f'<span class="tag">{word} ({count})</span>' for word, count in sorted_key_words[:10])
    remaining_key_words = "".join(f'<span class="tag">{word} ({count})</span>' for word, count in sorted_key_words[10:])

    # Write the HTML file with overview
    with open("document_review_report.html", "w") as f:
        f.write(html_content.format(
            file_name=file_name,  # Pass the file name
            word_count=overview['word_count'],
            page_count=overview['page_count'],
            top_key_words=top_key_words,
            remaining_key_words=remaining_key_words,
            total_errors=overview['total_errors'],
            error_types=error_types,
            error_sections=error_sections
        ))




# Function to analyze the document
def analyze_document(doc):
    word_count = 0
    key_word_counter = Counter()
    
    # Count total words and collect keywords
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            words = text.split()
            word_count += len(words)
            key_word_counter.update(word.lower() for word in words if word.lower() not in STOP_WORDS)

    # Calculate page count (assuming each section represents a page)
    page_count = len(doc.sections)
    
    # Get the key words as a dictionary with counts
    key_words = dict(key_word_counter)  # This will now be a dictionary
    
    # Total errors (initially set to 0, will be updated later)
    total_errors = 0  

    return {
        'word_count': word_count,
        'page_count': page_count,
        'key_words': key_words,  # This will now be a dictionary
        'total_errors': total_errors
    }


# Check basic formatting (Font, Font Size, Spacing)
def check_basic_formatting(doc, errors):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()  # Store trimmed text for further checks
        
        if text:  # Check if the sentence is not empty
            # Check for headings
            if not paragraph.style.name.startswith("Heading"):  # Adjust based on your heading styles
                for run in paragraph.runs:
                    # Check font name
                    if run.font.name != 'Times New Roman':
                        errors.append((text, f"Incorrect font: '{run.font.name}' found. Expected 'Times New Roman'."))
                    
                    # Check font size only if the text is longer than a single word or a short sentence
                    if len(text.split()) > 1:  # Adjust the number as needed for what you consider "short"
                        if run.font.size != Pt(12):
                            errors.append((text, f"Incorrect font size: '{run.font.size}' found. Expected 12pt."))
            else:
                # If it's a heading, skip font size and alignment checks
                continue

    # Check for spacing and justification (assuming body text is 'Normal' style)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()  # Store trimmed text for further checks
        
        if text:  # Check if the sentence is not empty
            if paragraph.style.name == 'Normal':
                if paragraph.paragraph_format.line_spacing != 1.5:
                    errors.append((text, f"Incorrect line spacing. Expected 1.5."))
                # Check if alignment is justified
                if paragraph.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    errors.append((text, "Incorrect justification. Expected justified alignment."))

# Check if name and roll number are present
def check_name_and_roll(doc, errors):
    # Assuming name and roll number are at the beginning or end
    text = "\n".join([p.text for p in doc.paragraphs])
    if "Name" not in text or "Roll Number" not in text:
        errors.append(("Entire Document", "Name or Roll Number is missing."))

# Check content structure (title, executive summary, company background)
def check_content_structure(doc, errors):
    sections = {
        "Title": False,
        "Executive Summary": False,
        "Company Background": False,
        "Problem Definition": False,
        "Objectives": False,
        "Data Collection": False,
        "Timelines": False,
        "Expected Outcomes": False,
        "Graphical Analysis": False
    }
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.lower()
        if "title" in text:
            sections["Title"] = True
        elif "executive summary" in text:
            sections["Executive Summary"] = True
        elif "company background" in text:
            sections["Company Background"] = True
        elif "problem definition" in text:
            sections["Problem Definition"] = True
        elif "objectives" in text:
            sections["Objectives"] = True
        elif "data collection" in text:
            sections["Data Collection"] = True
        elif "timelines" in text:
            sections["Timelines"] = True
        elif "expected outcomes" in text:
            sections["Expected Outcomes"] = True
        elif "graphical analysis" in text:
            sections["Graphical Analysis"] = True
    
    for section, present in sections.items():
        if not present:
            errors.append(("Entire Document", f"Section missing: {section}"))

# Check active/passive voice using NLP
def check_passive_voice(doc, errors):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:  # Only process non-empty sentences
            doc_nlp = nlp(text)
            for sent in doc_nlp.sents:
                if any(token.dep_ == "nsubjpass" for token in sent):
                    errors.append((paragraph.text, f"Passive voice detected: '{sent}'."))

# Check for page numbers and figure/table labels
def check_page_numbering(doc, errors):
    if doc.sections[0].footer.paragraphs[0].text.strip() == "":
        errors.append(("Entire Document", "Page numbers missing in the document."))
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Check if the sentence is not empty
            text = paragraph.text.lower()
            if "figure" in text or "table" in text:
                if not any(num in text for num in ['1', '2', '3', '4', '5']):
                    errors.append((paragraph.text, "Figure/Table label missing or incorrect."))

# Run all checks and generate HTML report
def run_checks(file_path):
    doc = load_document(file_path)
    errors = []
    
    # Analyze the document for overview
    overview = analyze_document(doc)
    
    print("Checking basic formatting...")
    check_basic_formatting(doc, errors)
    
    print("Checking name and roll number...")
    check_name_and_roll(doc, errors)
    
    print("Checking content structure...")
    check_content_structure(doc, errors)
    
    print("Checking for passive voice...")
    check_passive_voice(doc, errors)
    
    print("Checking page numbering and labels...")
    check_page_numbering(doc, errors)
    
    # Update the total errors count
    overview['total_errors'] = len(errors)
    
    # Generate HTML report with overview, passing the file name
    generate_html_output(overview, errors, file_name=file_path)

# Example usage
if __name__ == "__main__":
    # Change this to your Word file's path
    file_path = 'proposal.docx'
    run_checks(file_path)

